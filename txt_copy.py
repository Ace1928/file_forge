#!/usr/bin/env python3
"""
Text File Extractor

A sophisticated utility that finds and extracts text from various document types,
renames them using AI content analysis, and organizes them in a destination folder.

Key Features:
- Interactive CLI mode for real-time directory selection
- Support for multiple document types (txt, md, py, pdf, docx, etc.)
- AI-powered naming using Ollama's LLM models
- Original file type preservation in filename
- Deduplication using content hashing
- Parallel processing with dynamic adjustment based on system load
- Rich progress display and logging

Usage:
    python txt_copy.py                      # Run in interactive CLI mode
    python txt_copy.py -s /path -d /dest    # Run with command-line arguments
"""

import os
import sys
import shutil
import mimetypes
import tempfile
import time
import json
import hashlib
import argparse
import logging
import concurrent.futures
import psutil
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Set, Union, TypeVar, Any, cast
from rich.console import Console
from rich.prompt import Prompt, Confirm
from rich.panel import Panel
from rich.markdown import Markdown
from rich.table import Table
from rich.progress import (
        Progress, TaskID, TextColumn, BarColumn, 
        TimeElapsedColumn, TimeRemainingColumn, SpinnerColumn
)

# Document extraction imports - making them standard
import PyPDF2
import docx
import requests

# Configure logging
logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
                logging.FileHandler("text_extractor.log"),
                logging.StreamHandler()
        ]
)
logger = logging.getLogger(__name__)

# Initialize rich console for nice output
console = Console()

# Type definitions
PathLike = Union[str, Path]
FileContent = str
ContentHash = str
T = TypeVar('T')
# Update type definition for better future typing
ProcessResult = Tuple[Optional[Path], bool]
# Define proper type for futures
FutureResult = concurrent.futures.Future[Optional[Path]]

# Default Ollama API URLs
OLLAMA_BASE_URL = os.environ.get("OLLAMA_BASE_URL", "http://0.0.0.0:11434")
OLLAMA_GENERATE_URL = f"{OLLAMA_BASE_URL}/api/generate"
OLLAMA_CHAT_URL = f"{OLLAMA_BASE_URL}/api/chat"
OLLAMA_EMBEDDINGS_URL = f"{OLLAMA_BASE_URL}/api/embed"

# Default LLM model - can be overridden via environment variable
default_llm_model = os.environ.get("OLLAMA_MODEL", "deepseek-coder:1.5b")

# File type constants
TEXT_EXTENSIONS: Set[str] = {
        '.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', 
        '.csv', '.yml', '.yaml', '.sh', '.java', '.c', '.cpp', '.h', 
        '.rb', '.pl', '.php', '.ts', '.jsx', '.tsx', '.conf', '.ini',
        '.sql', '.rs', '.go', '.lua', '.r', '.swift'
}

# Mapping file extensions to type names for filename suffixes
FILE_TYPE_NAMES: Dict[str, str] = {
        '.txt': 'text',
        '.md': 'markdown',
        '.py': 'python',
        '.js': 'javascript',
        '.html': 'html',
        '.css': 'css',
        '.json': 'json',
        '.xml': 'xml',
        '.csv': 'csv',
        '.yml': 'yaml',
        '.yaml': 'yaml',
        '.pdf': 'pdf',
        '.docx': 'word',
        '.doc': 'word',
        '.sh': 'shell',
        '.java': 'java',
        '.c': 'c',
        '.cpp': 'cpp',
        '.h': 'header',
        '.ts': 'typescript',
        '.sql': 'sql',
        '.rs': 'rust',
        '.go': 'go',
}

# Add more specific types for Ollama API responses and requests
OllamaRequestDict = Dict[str, Union[str, bool, Dict[str, Any], List[Dict[str, Any]]]]
OllamaResponseDict = Dict[str, Any]

def check_ollama_available() -> bool:
    """
    Check if Ollama API is available
    
    Returns:
        bool: True if Ollama is responding, False otherwise
    """
    try:
        response = requests.get(f"{OLLAMA_BASE_URL}/api/version", timeout=2)
        if response.status_code == 200:
            version = response.json().get('version')
            logger.info(f"Ollama API available, version: {version}")
            return True
        else:
            logger.warning(f"Ollama API returned status code: {response.status_code}")
            return False
    except requests.exceptions.RequestException as e:
        logger.warning(f"Ollama API not available: {e}")
        return False

def get_optimal_workers() -> int:
    """
    Calculate optimal number of worker threads based on system resources.
    
    Returns:
        int: Recommended number of worker threads
    """
    cpu_count = os.cpu_count() or 4
    cpu_usage = psutil.cpu_percent(interval=0.5) / 100
    memory_usage = psutil.virtual_memory().percent / 100
    
    # Base worker count on CPU cores, adjusted for current load
    base_workers = max(1, cpu_count)
    load_factor = 1 - (0.5 * cpu_usage + 0.5 * memory_usage)
    
    # Adjust workers based on load (min 1, max cpu_count*2)
    workers = max(1, min(cpu_count * 2, int(base_workers * load_factor * 2)))
    
    logger.debug(f"System stats: CPU: {cpu_count} cores at {cpu_usage:.1%} usage, "
                 f"Memory: {memory_usage:.1%} used, Workers: {workers}")
    
    return workers


def monitor_system_resources() -> Dict[str, float]:
    """
    Monitor system resources to inform processing decisions.
    
    Returns:
        Dict[str, float]: Dictionary with resource usage information
    """
    return {
        'cpu_percent': psutil.cpu_percent(interval=0.5),
        'memory_percent': psutil.virtual_memory().percent,
        'disk_percent': psutil.disk_usage('/').percent
    }


def is_text_file(file_path: Path) -> bool:
    """
    Determine if a file is a text file based on extension, mimetype or content analysis.
    
    Args:
        file_path: Path to the file to check
            
    Returns:
        True if the file is a text file, False otherwise
    """
    # Skip files over 10MB by default
    if file_path.stat().st_size > 10 * 1024 * 1024:
        logger.debug(f"Skipping large file {file_path} ({file_path.stat().st_size / 1024 / 1024:.2f} MB)")
        return False
            
    # Check by extension first for common text file types
    if file_path.suffix.lower() in TEXT_EXTENSIONS:
        return True
    
    # Check for PDF, DOCX
    if file_path.suffix.lower() == '.pdf':
        return True
    if file_path.suffix.lower() in ('.docx', '.doc'):
        return True
    
    # Try mime type
    mime_type, _ = mimetypes.guess_type(str(file_path))
    if mime_type and mime_type.startswith('text/'):
        return True
    
    # Try reading the file as text as a last resort
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            f.read(4096)  # Read a chunk to check if it's readable as text
        return True
    except UnicodeDecodeError:
        return False
    except Exception as e:
        logger.debug(f"Error checking if file is text: {e}")
        return False


def extract_text_from_pdf(file_path: Path) -> Optional[str]:
    """
    Extract text content from a PDF file.
    
    Args:
        file_path: Path to the PDF file
            
    Returns:
        Extracted text content or None if extraction failed
    """
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            text: List[str] = []
            temp_path = Path(temp_dir) / "temp_pdf_extract.txt"
            
            # Use shutil for copying if needed
            if not temp_path.parent.exists():
                os.makedirs(temp_path.parent)
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages):
                    content = page.extract_text()
                    if content:
                        text.append(content)
                        
                        # Write content to temp file periodically to avoid memory issues
                        if page_num % 10 == 0 and text:
                            with open(temp_path, 'a', encoding='utf-8') as tmp:
                                tmp.write("\n\n".join(text))
                            text = []
            
            # Write any remaining content
            if text:
                with open(temp_path, 'a', encoding='utf-8') as tmp:
                    tmp.write("\n\n".join(text))
            
            # Read complete content
            if temp_path.exists() and temp_path.stat().st_size > 0:
                with open(temp_path, 'r', encoding='utf-8') as f:
                    return f.read()
            elif text:
                return "\n\n".join(text)
            else:
                return None
    except Exception as e:
        logger.error(f"Failed to extract text from PDF {file_path}: {e}")
        return None


def extract_text_from_docx(file_path: Path) -> Optional[str]:
    """
    Extract text content from a Word document.
    
    Args:
        file_path: Path to the Word document
            
    Returns:
        Extracted text content or None if extraction failed
    """
    try:
        # Convert Path to string for docx.Document
        doc = docx.Document(str(file_path))
        text: List[str] = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return "\n".join(text) if text else None
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
        return None


def extract_text_from_json(file_path: Path) -> Optional[str]:
    """
    Extract text content from a JSON file with proper formatting.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data: Dict[str, Any] = json.load(f)  # Provide explicit type
        return json.dumps(data, indent=2)
    except Exception as e:
        logger.error(f"Failed to extract text from JSON {file_path}: {e}")
        return extract_plain_text(file_path)


def extract_plain_text(file_path: Path) -> Optional[str]:
    """
    Extract plain text content from a file.
    
    Args:
        file_path: Path to the text file
            
    Returns:
        Text content or None if extraction failed
    """
    encoding_attempts: List[str] = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encoding_attempts:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            return None
    
    logger.error(f"Failed to extract text from {file_path} with all attempted encodings")
    return None


def extract_text_content(file_path: Path) -> Optional[str]:
    """
    Extract text content from various file types based on extension.
    
    Args:
        file_path: Path to the file
            
    Returns:
        Extracted text content or None if extraction failed
    """
    suffix = file_path.suffix.lower()
    
    if suffix == '.pdf':
        return extract_text_from_pdf(file_path)
    elif suffix in ('.docx', '.doc'):
        return extract_text_from_docx(file_path)
    elif suffix == '.json':
        return extract_text_from_json(file_path)
    else:
        return extract_plain_text(file_path)


def generate_name_with_ollama(content: str, max_retries: int = 3) -> Optional[str]:
    """
    Generate a concise, descriptive filename (without extension) for the given text content
    using Ollama's API with the configured LLM model.
    
    Args:
        content: Text content to generate a name for
        max_retries: Maximum number of retries on failure
            
    Returns:
        Generated name or None if generation failed
    """
    # Check if content is empty
    if not content or len(content.strip()) < 10:
        logger.warning("Content is empty or too short for naming.")
        return None

    # Truncate content for the prompt (limited to 1000 chars to avoid large requests)
    truncated = content[:500] + ("..." if len(content) > 500 else "")
    
    # Try different API endpoints based on availability and success
    # First try the generate endpoint (simpler, faster)
    for attempt in range(max_retries):
        try:
            # Using the generate API endpoint first
            generate_data: OllamaRequestDict = {
                "model": default_llm_model,
                "prompt": (
                    "Generate a concise, descriptive filename (without extension) for this text content. "
                    "Maximum 30 characters, use only alphanumeric characters, hyphens, and underscores. "
                    "No spaces or special characters. The filename should clearly describe this content:\n\n"
                    f"{truncated}\n\n"
                    "Filename (without extension):"
                ),
                "stream": False,
            }
            
            response = requests.post(
                OLLAMA_GENERATE_URL,
                json=generate_data,
            )
            
            if response.status_code == 200:
                result = response.json()
                name = result.get("response", "").strip()
                logger.debug(f"Ollama API generate response: {name}")
                
                # If we got a good response, clean and return it
                if name and len(name) > 2:
                    name = ''.join(c if c.isalnum() or c in '-_' else '_' for c in name)
                    name = name[:30].strip('_')
                    if name and name not in ['.', '..']:
                        logger.info(f"Generated name: {name}")
                        return name
            
            # If we're here, we need to try again or fall back to the chat API
            logger.warning(f"Ollama generate API attempt {attempt+1} failed, will try again")
            time.sleep(1)
            
        except requests.exceptions.RequestException as e:
            logger.warning(f"Ollama generate API error: {e}")
            time.sleep(1)
    
    # Fall back to the chat API if generate didn't work
    for attempt in range(max_retries):
        try:
            # Using the chat API as fallback
            chat_data: OllamaRequestDict = {
                "model": default_llm_model,
                "messages": [
                    {
                        "role": "system", 
                        "content": (
                            "You are a file naming assistant. Generate a concise filename "
                            "(no extension) for the content. Use only alphanumeric characters, "
                            "hyphens and underscores. Maximum 30 characters."
                        )
                    },
                    {
                        "role": "user", 
                        "content": f"Create a descriptive filename for this content:\n{truncated}"
                    }
                ],
                "stream": False
            }
            
            response = requests.post(
                OLLAMA_CHAT_URL,
                json=chat_data,
                timeout=10
            )
            
            if response.status_code == 200:
                result = response.json()
                name = result.get("message", {}).get("content", "").strip()
                logger.debug(f"Ollama API chat response: {name}")
                
                if name and len(name) > 2:
                    name = ''.join(c if c.isalnum() or c in '-_' else '_' for c in name)
                    name = name[:30].strip('_')
                    if name and name not in ['.', '..']:
                        logger.info(f"Generated name (chat): {name}")
                        return name
            
            logger.warning(f"Ollama chat API attempt {attempt+1} failed")
            time.sleep(1)
            
        except requests.exceptions.RequestException as e:
            logger.warning(f"Ollama chat API error: {e}")
            time.sleep(1)
    
    # Final fallback: hash-based name
    fallback_name = f"text_{hashlib.md5(content[:500].encode()).hexdigest()[:10]}"
    logger.info(f"Using fallback filename: {fallback_name}")
    return fallback_name


def get_content_embedding(content: str, max_length: int = 2000) -> Optional[List[float]]:
    """
    Get embedding vector for text content using Ollama API
    
    Args:
        content: Text to get embedding for
        max_length: Maximum length of content to embed
        
    Returns:
        List of float values representing the embedding, or None if failed
    """
    # Truncate content to avoid large requests
    truncated = content[:max_length]
    
    try:
        response = requests.post(
            OLLAMA_EMBEDDINGS_URL,
            json=cast(Dict[str, Any], {
                "model": "all-minilm",  # Use embedding-specific model
                "prompt": truncated
            }),
            timeout=5
        )
        
        if response.status_code == 200:
            result = response.json()
            embedding = result.get("embedding")
            if embedding and isinstance(embedding, list):
                # Cast the embedding to List[float] to fix the partial unknown warning
                return cast(List[float], embedding)
        
        logger.warning(f"Failed to get embedding: {response.status_code} - {response.text}")
        return None
        
    except requests.exceptions.RequestException as e:
        logger.warning(f"Error getting embedding: {e}")
        return None


def find_text_files(source_dir: Path, progress: Optional[Progress] = None, task_id: Optional[TaskID] = None) -> List[Path]:
    """
    Find all text files in the source directory recursively.
    
    Args:
        source_dir: Directory to search
        progress: Rich progress instance for display
        task_id: Task ID for the progress bar
            
    Returns:
        List of paths to text files
    """
    text_files: List[Path] = []
    total_files = 0
    
    # Count files first for progress tracking
    for root, _, files in os.walk(source_dir):
        total_files += len(files)
    
    if progress and task_id:
        progress.update(task_id, total=total_files)
    
    processed = 0
    
    for root, _, files in os.walk(source_dir):
        for file in files:
            file_path = Path(root) / file
            try:
                if is_text_file(file_path):
                    text_files.append(file_path)
            except Exception as e:
                logger.error(f"Error processing file {file_path}: {e}")
            
            processed += 1
            if progress and task_id:
                progress.update(task_id, completed=processed)
    
    return text_files


def process_file(file_path: Path, dest_dir: Path, processed_hashes: Set[str]) -> Optional[Path]:
    """
    Process a single text file and copy it to the destination with a new name.
    
    Args:
        file_path: Path to the file to process
        dest_dir: Destination directory
        processed_hashes: Set of already processed content hashes
            
    Returns:
        Path to the created file or None if processing failed
    """
    try:
        # Extract content
        content = extract_text_content(file_path)
        if not content or len(content.strip()) < 10:  # Skip empty or very small files
            logger.info(f"Skipping empty or very small file: {file_path}")
            return None
            
        # Calculate content hash to avoid duplicates
        content_hash = hashlib.md5(content.encode()).hexdigest()
        if content_hash in processed_hashes:
            logger.info(f"Skipping duplicate content from {file_path}")
            return None
            
        processed_hashes.add(content_hash)
        
        # Get file type for name suffix
        file_type = FILE_TYPE_NAMES.get(file_path.suffix.lower(), 'unknown')
        
        # Generate name
        base_name = generate_name_with_ollama(content)
        if not base_name:
            base_name = f"text_{content_hash[:10]}"
        
        # Append file type to name
        new_name = f"{base_name}_{file_type}"
        
        # Ensure unique filename
        dest_path = dest_dir / f"{new_name}.txt"
        counter = 1
        while dest_path.exists():
            dest_path = dest_dir / f"{new_name}_{counter}.txt"
            counter += 1
        
        # Write content to new file - using shutil to copy instead of direct write when possible
        if file_path.suffix.lower() == '.txt' and not content_hash in processed_hashes:
            # Use shutil for text files (direct copy)
            shutil.copy2(file_path, dest_path)
            
            # Add metadata header with original file info
            with open(dest_path, 'r+', encoding='utf-8') as f:
                original_content = f.read()
                f.seek(0)
                metadata = (
                    f"# Original File: {file_path}\n"
                    f"# Original Type: {file_type}\n"
                    f"# Extraction Date: {time.strftime('%Y-%m-%d %H:%M:%S')}\n"
                    f"# Content Hash: {content_hash}\n"
                    f"{'#' * 50}\n\n"
                )
                f.write(metadata + original_content)
        else:
            # Direct write for other types
            with open(dest_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            # Add metadata header with original file info
            with open(dest_path, 'r+', encoding='utf-8') as f:
                original_content = f.read()
                f.seek(0)
                metadata = (
                    f"# Original File: {file_path}\n"
                    f"# Original Type: {file_type}\n"
                    f"# Extraction Date: {time.strftime('%Y-%m-%d %H:%M:%S')}\n"
                    f"# Content Hash: {content_hash}\n"
                    f"{'#' * 50}\n\n"
                )
                f.write(metadata + original_content)
            
        return dest_path
    except Exception as e:
        logger.error(f"Failed to process {file_path}: {e}")
        return None


def browse_directory(start_path: Optional[Path] = None) -> Path:
    """
    Interactive directory browser.
    
    Args:
        start_path: Path to start browsing from
            
    Returns:
        Selected directory path
    """
    if start_path is None:
        start_path = Path.home()
    
    current_path = start_path.resolve()
    
    while True:
        console.clear()
        console.print(Panel(f"[bold blue]Directory Browser - Current: [/bold blue][yellow]{current_path}[/yellow]"))
        
        # List directories
        dirs = [d for d in current_path.iterdir() if d.is_dir()]
        dirs.sort()
        
        # Create table
        table = Table(show_header=True, header_style="bold")
        table.add_column("#", style="dim")
        table.add_column("Directory", style="green")
        
        # Add special options
        table.add_row("0", "[yellow].. (Parent Directory)[/yellow]")
        table.add_row("S", "[bold green]SELECT THIS DIRECTORY[/bold green]")
        
        # Add directories
        for i, d in enumerate(dirs, 1):
            table.add_row(str(i), d.name)
        
        console.print(table)
        
        # Get user choice
        choice = Prompt.ask("Enter your choice (number, S to select, or full path)")
        
        if choice.upper() == "S":
            return current_path
        elif choice == "0":
            current_path = current_path.parent
        elif choice.isdigit() and 1 <= int(choice) <= len(dirs):
            current_path = dirs[int(choice) - 1]
        elif Path(choice).exists():
            # Allow direct path entry
            current_path = Path(choice).resolve()
        else:
            console.print("[red]Invalid choice[/red]")
            time.sleep(1)


def show_welcome_banner() -> None:
    """Display a welcome banner with application information."""
    welcome_text = """
    # Text Extraction Utility
    
    A sophisticated tool to find and extract text from various document types,
    rename them using AI content analysis, and organize them in a destination folder.
    
    ## Features:
    - Extract text from TXT, MD, PY, PDF, DOCX, JSON and many more file types
    - AI-powered naming using Ollama's LLM models
    - Original file type preservation in filename
    - Deduplication using content hashing
    - Parallel processing for improved performance
    """
    
    console.print(Panel(Markdown(welcome_text), title="Welcome", border_style="green"))


def interactive_mode() -> int:
    """
    Run the application in interactive CLI mode.
    """
    show_welcome_banner()
    
    # Check if Ollama is available
    ollama_available = check_ollama_available()
    if not ollama_available:
        console.print("[yellow]Warning: Ollama API is not available. AI-powered naming will fall back to hash-based names.[/yellow]")
        if not Confirm.ask("Continue without Ollama API?"):
            console.print("[yellow]Operation cancelled by user[/yellow]")
            return 0
    
    # Source directory selection
    console.print("[bold]Source Directory Selection[/bold]")
    source_dir = browse_directory()
    console.print(f"Selected source directory: [green]{source_dir}[/green]")
    
    # Destination directory selection
    console.print("\n[bold]Destination Directory Selection[/bold]")
    use_default_dest = Confirm.ask("Use current directory as destination?")
    if use_default_dest:
        dest_dir = Path.cwd()
    else:
        dest_dir = browse_directory()
    console.print(f"Selected destination directory: [green]{dest_dir}[/green]")
    
    # Additional options
    console.print("\n[bold]Processing Options[/bold]")
    max_jobs = int(Prompt.ask("Number of parallel jobs", default=str(get_optimal_workers())))
    
    # Model selection if Ollama is available
    model = default_llm_model
    if ollama_available:
        try:
            # Try to get list of available models
            response = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=2)
            if response.status_code == 200:
                models = [m.get("name") for m in response.json().get("models", [])]
                if models:
                    console.print("\n[bold]Available models:[/bold]")
                    for i, model_name in enumerate(models, 1):
                        console.print(f"  {i}. {model_name}")
                    model_choice = Prompt.ask(
                        f"Select model (1-{len(models)}, or leave empty for default '{default_llm_model}')",
                        default=""
                    )
                    if model_choice.isdigit() and 1 <= int(model_choice) <= len(models):
                        model = models[int(model_choice) - 1]
                        console.print(f"Selected model: [green]{model}[/green]")
                        # Update default model (not using global keyword to avoid constant redefinition)
                        globals()["default_llm_model"] = model
        except Exception as e:
            logger.warning(f"Could not fetch model list: {e}")
    
    # Confirmation
    console.print("\n[bold]Configuration Summary:[/bold]")
    console.print(f"Source: [blue]{source_dir}[/blue]")
    console.print(f"Destination: [blue]{dest_dir}[/blue]")
    console.print(f"Parallel Jobs: [blue]{max_jobs}[/blue]")
    if ollama_available:
        console.print(f"AI Model: [blue]{model}[/blue]")
    
    if not Confirm.ask("\nProceed with extraction?"):
        console.print("[yellow]Operation cancelled by user[/yellow]")
        return 0
    
    # Run the actual processing
    return run_extraction(source_dir, dest_dir, max_jobs)


def run_extraction(source_dir: Path, dest_dir: Path, jobs: int) -> int:
    """
    Run the extraction process with the specified parameters.
    
    Args:
        source_dir: Source directory
        dest_dir: Destination directory
        jobs: Number of parallel jobs
            
    Returns:
        Exit code (0 for success)
    """
    if not source_dir.exists():
        console.print(f"[bold red]Source directory '{source_dir}' does not exist![/]")
        return 1
            
    if not dest_dir.exists():
        try:
            dest_dir.mkdir(parents=True)
            console.print(f"Created destination directory: [green]{dest_dir}[/green]")
        except Exception as e:
            console.print(f"[bold red]Failed to create destination directory: {e}[/]")
            return 1
    
    console.print(f"[bold green]Starting text file extraction[/]")
    console.print(f"Source: [blue]{source_dir}[/]")
    console.print(f"Destination: [blue]{dest_dir}[/]")
    
    processed_hashes: Set[str] = set()
    successful_copies = 0
    start_time = time.time()
    
    # Dynamic worker calculation based on system load
    initial_jobs = min(jobs, get_optimal_workers())
    console.print(f"Starting with [blue]{initial_jobs}[/] worker threads based on system resources")
    
    with Progress(
            SpinnerColumn(),
            TextColumn("[bold blue]{task.description}"),
            BarColumn(),
            TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
            TimeElapsedColumn(),
            TimeRemainingColumn(),
            console=console
    ) as progress:
        # Find all text files
        find_task = progress.add_task("[green]Finding text files...", total=None)
        text_files = find_text_files(source_dir, progress, find_task)
        progress.update(find_task, completed=progress.tasks[find_task].total)
        
        console.print(f"[green]Found {len(text_files)} text files[/]")
        
        # Process files with dynamic thread pool size adjustment
        process_task = progress.add_task("[yellow]Processing files...", total=len(text_files))
        
        # Track resources and adjust thread count periodically
        last_check_time = time.time()
        current_workers = initial_jobs
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=current_workers) as executor:
            # Properly type the future dictionary
            future_to_file: Dict[FutureResult, Path] = {}
            
            # Submit initial batch of tasks
            for file_path in text_files[:current_workers*2]:  # Start with a batch
                future = executor.submit(process_file, file_path, dest_dir, processed_hashes)
                future_to_file[future] = file_path
            
            remaining_files = text_files[current_workers*2:]
            
            # Process results and submit new tasks
            while future_to_file or remaining_files:
                # Check system resources every 5 seconds
                if time.time() - last_check_time > 5:
                    resources = monitor_system_resources()
                    optimal_workers = get_optimal_workers()
                    
                    # Adjust worker count if needed
                    if optimal_workers != current_workers:
                        console.print(f"Adjusting workers from {current_workers} to {optimal_workers} "
                                     f"(CPU: {resources['cpu_percent']}%, Memory: {resources['memory_percent']}%)")
                        current_workers = optimal_workers
                        # Note: ThreadPoolExecutor doesn't allow dynamic resizing, 
                        # but new tasks will use the updated worker count
                    
                    last_check_time = time.time()
                
                # Process completed tasks - properly handle typing
                done, _ = concurrent.futures.wait(
                    list(future_to_file.keys()),
                    timeout=0.5,
                    return_when=concurrent.futures.FIRST_COMPLETED
                )
                
                for future in done:
                    file_path = future_to_file.pop(future)
                    try:
                        result = future.result()
                        if result:
                            successful_copies += 1
                    except Exception as e:
                        logger.error(f"Error processing {file_path}: {e}")
                    
                    progress.update(process_task, advance=1)
                    
                    # Submit a new task if there are files left to process
                    if remaining_files:
                        next_file = remaining_files.pop(0)
                        new_future = executor.submit(process_file, next_file, dest_dir, processed_hashes)
                        future_to_file[new_future] = next_file
    
    elapsed = time.time() - start_time
    
    # Final summary
    console.print("\n[bold]Extraction Summary:[/]")
    console.print(f"Total files scanned: [blue]{len(text_files)}[/]")
    console.print(f"Files successfully copied: [green]{successful_copies}[/]")
    console.print(f"Duplicates skipped: [yellow]{len(text_files) - successful_copies}[/]")
    console.print(f"Total time: [blue]{elapsed:.2f} seconds[/]")
    console.print(f"Files per second: [blue]{len(text_files) / elapsed:.2f}[/]")
    console.print(f"[bold green]Done! Files copied to {dest_dir}[/]")
    
    return 0


def main() -> int:
    """
    Main entry point for the application.
    
    Returns:
        Exit code (0 for success)
    """
    parser = argparse.ArgumentParser(
            description='Extract and copy text from various file types with AI-generated names',
            formatter_class=argparse.RawDescriptionHelpFormatter,
            epilog="""
Examples:
    python txt_copy.py                      # Run in interactive CLI mode
    python txt_copy.py -s ~/Documents       # Process documents with interactive destination
    python txt_copy.py -s ~/src -d ~/dest   # Process with specified source and destination
    python txt_copy.py -j 8                 # Use 8 parallel jobs
            """
    )
    parser.add_argument('-s', '--source', type=str,
                       help='Source directory to copy from (interactive if not specified)')
    parser.add_argument('-d', '--dest', type=str,
                       help='Destination directory (interactive if not specified)')
    parser.add_argument('-j', '--jobs', type=int, default=get_optimal_workers(),
                       help=f'Number of parallel jobs (default: auto={get_optimal_workers()})')
    parser.add_argument('--interactive', action='store_true',
                       help='Force interactive mode even if source/dest are provided')
    parser.add_argument('--model', type=str, default=default_llm_model,
                       help=f'Ollama model to use for naming (default: {default_llm_model})')
    
    args = parser.parse_args()
    
    # Update model if specified (using globals instead of global keyword)
    if args.model:
        globals()["default_llm_model"] = args.model
    
    # Check if we need interactive mode
    if args.interactive or (args.source is None and args.dest is None):
        return interactive_mode()
    
    # If we have at least some arguments, use them
    source_dir = Path(args.source).expanduser().resolve() if args.source else None
    dest_dir = Path(args.dest).expanduser().resolve() if args.dest else None
    
    if source_dir is None:
        console.print("[bold]Source Directory Selection[/bold]")
        source_dir = browse_directory()
            
    if dest_dir is None:
        console.print("\n[bold]Destination Directory Selection[/bold]")
        use_default_dest = Confirm.ask("Use current directory as destination?")
        if use_default_dest:
            dest_dir = Path.cwd()
        else:
            dest_dir = browse_directory()
    
    return run_extraction(source_dir, dest_dir, args.jobs)


if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        console.print("\n[bold red]Operation cancelled by user[/]")
        sys.exit(1)
    except Exception as e:
        console.print(f"\n[bold red]An unexpected error occurred: {e}[/]")
        logger.exception("Unexpected error")
        sys.exit(1)